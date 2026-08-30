from __future__ import annotations

import hashlib
import io
import json
import zipfile
from pathlib import Path
from typing import Any

import pandas as pd
import pyarrow.parquet as parquet
from docx import Document
from PIL import Image
from pypdf import PdfReader

from .contracts import (
    AttachmentKind,
    EvidenceLocator,
    ParsedAttachment,
    TableData,
    TextEvidence,
)
from .security import PromptInjectionDetected, reject_prompt_injection


class FileParseError(ValueError):
    pass


_MIME = {
    ".csv": {"text/csv", "application/csv", "text/plain"},
    ".xls": {"application/vnd.ms-excel", "application/octet-stream"},
    ".xlsx": {"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/zip"},
    ".parquet": {"application/vnd.apache.parquet", "application/octet-stream"},
    ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/zip"},
    ".pdf": {"application/pdf"},
    ".txt": {"text/plain"},
    ".md": {"text/plain", "text/markdown"},
    ".png": {"image/png"},
    ".jpg": {"image/jpeg"},
    ".jpeg": {"image/jpeg"},
    ".webp": {"image/webp"},
}

_MAX_OFFICE_MEMBERS = 10_000
_MAX_OFFICE_UNCOMPRESSED_BYTES = 128 * 1024 * 1024
_MAX_OFFICE_COMPRESSION_RATIO = 200


def _validate_office_archive(archive: zipfile.ZipFile) -> set[str]:
    members = archive.infolist()
    if len(members) > _MAX_OFFICE_MEMBERS:
        raise FileParseError("OFFICE_ARCHIVE_MEMBER_LIMIT_EXCEEDED")
    total_uncompressed = sum(item.file_size for item in members)
    total_compressed = sum(item.compress_size for item in members)
    if total_uncompressed > _MAX_OFFICE_UNCOMPRESSED_BYTES:
        raise FileParseError("OFFICE_ARCHIVE_EXPANDED_SIZE_EXCEEDED")
    if total_uncompressed > 0 and total_uncompressed / max(1, total_compressed) > _MAX_OFFICE_COMPRESSION_RATIO:
        raise FileParseError("OFFICE_ARCHIVE_COMPRESSION_RATIO_EXCEEDED")
    if any(item.file_size > _MAX_OFFICE_UNCOMPRESSED_BYTES for item in members):
        raise FileParseError("OFFICE_ARCHIVE_MEMBER_SIZE_EXCEEDED")
    return {item.filename for item in members}


def _validate_signature(extension: str, declared_mime: str, data: bytes) -> None:
    if extension not in _MIME or declared_mime.lower() not in _MIME[extension]:
        raise FileParseError("MIME_EXTENSION_MISMATCH")
    if not data:
        raise FileParseError("EMPTY_FILE")
    if extension == ".pdf" and not data.startswith(b"%PDF-"):
        raise FileParseError("FILE_SIGNATURE_MISMATCH")
    if extension == ".xls" and not data.startswith(bytes.fromhex("D0CF11E0A1B11AE1")):
        raise FileParseError("FILE_SIGNATURE_MISMATCH")
    if extension == ".parquet" and not (data.startswith(b"PAR1") and data.endswith(b"PAR1")):
        raise FileParseError("FILE_SIGNATURE_MISMATCH")
    if extension in {".xlsx", ".docx"}:
        if not data.startswith(b"PK"):
            raise FileParseError("FILE_SIGNATURE_MISMATCH")
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as archive:
                names = _validate_office_archive(archive)
        except (zipfile.BadZipFile, OSError) as exc:
            raise FileParseError("INVALID_OFFICE_ARCHIVE") from exc
        required = "xl/workbook.xml" if extension == ".xlsx" else "word/document.xml"
        if required not in names:
            raise FileParseError("OFFICE_STRUCTURE_MISMATCH")
    if extension == ".png" and not data.startswith(b"\x89PNG\r\n\x1a\n"):
        raise FileParseError("FILE_SIGNATURE_MISMATCH")
    if extension in {".jpg", ".jpeg"} and not data.startswith(b"\xff\xd8\xff"):
        raise FileParseError("FILE_SIGNATURE_MISMATCH")
    if extension == ".webp" and not (data.startswith(b"RIFF") and data[8:12] == b"WEBP"):
        raise FileParseError("FILE_SIGNATURE_MISMATCH")
    if extension in {".csv", ".txt", ".md"}:
        if b"\x00" in data[:8192]:
            raise FileParseError("BINARY_TEXT_FILE")
        try:
            data.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise FileParseError("TEXT_MUST_BE_UTF8") from exc


def _json_value(value: Any) -> Any:
    if pd.isna(value):
        return None
    if hasattr(value, "isoformat"):
        return value.isoformat()
    if hasattr(value, "item"):
        return value.item()
    return value


def _table(name: str, frame: pd.DataFrame, *, max_rows: int) -> TableData:
    if len(frame) > max_rows:
        raise FileParseError("FILE_ROW_LIMIT_EXCEEDED")
    columns = tuple(str(value) for value in frame.columns)
    rows = tuple(
        {column: _json_value(value) for column, value in zip(columns, row)}
        for row in frame.itertuples(index=False, name=None)
    )
    reject_prompt_injection(
        str(value) for row in rows for value in row.values() if isinstance(value, str)
    )
    return TableData(name=name, columns=columns, rows=rows, row_count=len(rows))


def _result_signature(
    *, kind: AttachmentKind, tables: tuple[TableData, ...], evidence: tuple[TextEvidence, ...], page_count: int
) -> str:
    payload = {
        "kind": kind.value,
        "tables": [
            {"name": table.name, "columns": table.columns, "rows": table.rows, "row_count": table.row_count}
            for table in tables
        ],
        "text_evidence": [
            {"text": item.text, "locator": item.locator.__dict__} for item in evidence
        ],
        "page_count": page_count,
    }
    # Preserve canonical_sha256's exact wire format without materialising a
    # second full JSON byte string for large structured attachments.
    digest = hashlib.sha256()
    encoder = json.JSONEncoder(
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        default=str,
    )
    for chunk in encoder.iterencode(payload):
        digest.update(chunk.encode("utf-8"))
    return digest.hexdigest()


def parse_attachment(
    filename: str,
    declared_mime: str,
    data: bytes,
    *,
    max_rows: int = 100_000,
) -> ParsedAttachment:
    extension = Path(filename).suffix.lower()
    _validate_signature(extension, declared_mime, data)
    source = io.BytesIO(data)
    tables: tuple[TableData, ...] = ()
    evidence: tuple[TextEvidence, ...] = ()
    page_count = 0
    requires_vision = False

    if extension == ".csv":
        tables = (_table("data", pd.read_csv(source, nrows=max_rows + 1), max_rows=max_rows),)
        kind = AttachmentKind.STRUCTURED
    elif extension in {".xls", ".xlsx"}:
        collected: list[TableData] = []
        remaining_rows = max_rows
        with pd.ExcelFile(source) as workbook:
            if not workbook.sheet_names:
                raise FileParseError("EMPTY_WORKBOOK")
            for name in workbook.sheet_names:
                # Read at most one row beyond the remaining global allowance so
                # a many-sheet workbook cannot allocate N * max_rows objects
                # before the aggregate limit is checked.
                frame = pd.read_excel(workbook, sheet_name=name, nrows=remaining_rows + 1)
                if len(frame.columns) == 0:
                    # Workbooks commonly keep blank template tabs. They are
                    # not data tables and must not invalidate populated tabs.
                    continue
                table = _table(str(name), frame, max_rows=remaining_rows)
                collected.append(table)
                remaining_rows -= table.row_count
        if not collected:
            raise FileParseError("EMPTY_WORKBOOK")
        tables = tuple(collected)
        kind = AttachmentKind.STRUCTURED
    elif extension == ".parquet":
        if parquet.ParquetFile(source).metadata.num_rows > max_rows:
            raise FileParseError("FILE_ROW_LIMIT_EXCEEDED")
        source.seek(0)
        tables = (_table("data", pd.read_parquet(source), max_rows=max_rows),)
        kind = AttachmentKind.STRUCTURED
    elif extension == ".docx":
        document = Document(source)
        collected: list[TextEvidence] = []
        for index, paragraph in enumerate(document.paragraphs, start=1):
            if paragraph.text.strip():
                collected.append(TextEvidence(paragraph.text.strip(), EvidenceLocator("docx.paragraph", paragraph=index)))
        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                for column_index, cell in enumerate(row.cells, start=1):
                    if cell.text.strip():
                        collected.append(TextEvidence(
                            cell.text.strip(),
                            EvidenceLocator("docx.table_cell", table=table_index, row=row_index, column=str(column_index)),
                        ))
        reject_prompt_injection(item.text for item in collected)
        evidence = tuple(collected)
        kind = AttachmentKind.DOCUMENT
    elif extension == ".pdf":
        reader = PdfReader(source)
        page_count = len(reader.pages)
        collected = []
        blank_pages = 0
        for page_index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                collected.append(TextEvidence(text, EvidenceLocator("pdf.page", page=page_index)))
            else:
                blank_pages += 1
        reject_prompt_injection(item.text for item in collected)
        evidence = tuple(collected)
        requires_vision = blank_pages > 0
        kind = AttachmentKind.SCANNED_PDF if page_count and blank_pages == page_count else AttachmentKind.DOCUMENT
    elif extension in {".txt", ".md"}:
        text = data.decode("utf-8-sig")
        reject_prompt_injection((text,))
        evidence = (TextEvidence(text, EvidenceLocator("text.document")),)
        kind = AttachmentKind.DOCUMENT
    else:
        try:
            with Image.open(source) as image:
                image.verify()
        except Exception as exc:
            raise FileParseError("INVALID_IMAGE") from exc
        kind = AttachmentKind.IMAGE
        requires_vision = True

    signature = _result_signature(kind=kind, tables=tables, evidence=evidence, page_count=page_count)
    return ParsedAttachment(
        filename=Path(filename).name,
        extension=extension,
        mime_type=declared_mime.lower(),
        file_sha256=hashlib.sha256(data).hexdigest(),
        kind=kind,
        tables=tables,
        text_evidence=evidence,
        page_count=page_count,
        requires_vision=requires_vision,
        result_signature=signature,
    )


__all__ = ["FileParseError", "PromptInjectionDetected", "parse_attachment"]
